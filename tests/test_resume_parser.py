"""Resume text extraction."""

import io

import fitz
import pytest
from docx import Document

from app.config import settings
from app.resumes import parser

SAMPLE = (
    "Jane Doe\n"
    "Senior Software Engineer\n"
    "6 years of experience building distributed systems.\n"
    "Python, React, PostgreSQL, Docker, Kubernetes\n"
    "Bengaluru, India\n"
    "Prefers remote work\n"
)


def make_pdf(text: str = SAMPLE) -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    data = doc.tobytes()
    doc.close()
    return data


def make_docx(text: str = SAMPLE) -> bytes:
    document = Document()
    for line in text.split("\n"):
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def test_extracts_text_from_pdf() -> None:
    text = await parser.extract_text(make_pdf(), "resume.pdf")
    assert "Jane Doe" in text
    assert "PostgreSQL" in text


async def test_extracts_text_from_docx() -> None:
    text = await parser.extract_text(make_docx(), "resume.docx")
    assert "Jane Doe" in text
    assert "Kubernetes" in text


async def test_extracts_docx_table_cells() -> None:
    """Skills often live in a table, which `paragraphs` alone would skip."""
    document = Document()
    document.add_paragraph("Jane Doe, Senior Software Engineer in Bengaluru India")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Languages"
    table.rows[0].cells[1].text = "Python, Rust, Elixir"
    buffer = io.BytesIO()
    document.save(buffer)

    text = await parser.extract_text(buffer.getvalue(), "resume.docx")
    assert "Elixir" in text


async def test_extension_is_matched_case_insensitively() -> None:
    text = await parser.extract_text(make_pdf(), "RESUME.PDF")
    assert "Jane Doe" in text


@pytest.mark.parametrize("filename", ["resume.txt", "resume.rtf", "resume", "a.pages"])
async def test_unsupported_extensions_are_rejected(filename: str) -> None:
    with pytest.raises(ValueError, match="Unsupported file type"):
        await parser.extract_text(make_pdf(), filename)


async def test_legacy_doc_gets_its_own_message() -> None:
    """.doc is OLE2, not a renamed .docx — a generic error would send people
    in circles renaming the file."""
    with pytest.raises(ValueError, match="Legacy .doc"):
        await parser.extract_text(make_pdf(), "resume.doc")


async def test_empty_upload_is_rejected() -> None:
    with pytest.raises(ValueError, match="empty"):
        await parser.extract_text(b"", "resume.pdf")


async def test_oversized_upload_is_rejected_before_parsing() -> None:
    oversized = b"x" * (settings.max_resume_size_bytes + 1)
    with pytest.raises(ValueError, match="too large"):
        await parser.extract_text(oversized, "resume.pdf")


async def test_text_free_pdf_is_rejected() -> None:
    """An image-only scan parses fine but yields nothing useful."""
    doc = fitz.open()
    doc.new_page()
    data = doc.tobytes()
    doc.close()

    with pytest.raises(ValueError, match="Almost no text"):
        await parser.extract_text(data, "scan.pdf")


async def test_corrupt_pdf_raises_a_readable_error() -> None:
    with pytest.raises(ValueError):
        await parser.extract_text(b"%PDF-1.4 not really a pdf", "broken.pdf")


async def test_corrupt_docx_raises_a_readable_error() -> None:
    with pytest.raises(ValueError, match="could not be read"):
        await parser.extract_text(b"PK\x03\x04 garbage", "broken.docx")


def test_normalize_collapses_whitespace_noise() -> None:
    messy = "Jane   Doe\r\n\r\n\r\n\r\n  Senior  Engineer   \r\n"
    assert parser.normalize_text(messy) == "Jane Doe\n\nSenior Engineer"
