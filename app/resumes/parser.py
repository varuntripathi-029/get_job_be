"""Text extraction from uploaded resume files.

Parsing runs on user-supplied bytes, so every failure mode here is a 4xx rather
than a crash: unsupported type, oversized file, encrypted PDF, or a scan with no
text layer all raise ValueError with a message safe to show the uploader.
"""

from __future__ import annotations

import io
import logging
import re

import fitz  # PyMuPDF
from docx import Document

from app.config import settings

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = (".pdf", ".docx")
# A resume shorter than this is almost always a failed text layer (a scan or an
# image-only export) rather than a genuinely terse candidate.
MIN_TEXT_CHARS = 50

_WHITESPACE_RUN = re.compile(r"[ \t ]+")
_BLANK_LINES = re.compile(r"\n{3,}")


def _extension(filename: str) -> str:
    _, _, ext = filename.lower().rpartition(".")
    return f".{ext}" if ext else ""


def normalize_text(raw: str) -> str:
    """Collapse the whitespace noise that PDF extraction leaves behind."""
    text = raw.replace("\r\n", "\n").replace("\r", "\n")
    text = _WHITESPACE_RUN.sub(" ", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    return _BLANK_LINES.sub("\n\n", text).strip()


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            if doc.needs_pass:
                raise ValueError(
                    "This PDF is password-protected. Upload an unlocked copy."
                )
            return "\n".join(page.get_text() for page in doc)
    except ValueError:
        raise
    except Exception as exc:  # noqa: BLE001 — MuPDF raises its own error types
        logger.warning("PDF parse failed: %s", exc)
        raise ValueError("This PDF could not be read. It may be corrupt.") from exc


def _extract_docx(file_bytes: bytes) -> str:
    try:
        document = Document(io.BytesIO(file_bytes))
    except Exception as exc:  # noqa: BLE001 — python-docx raises PackageNotFoundError
        logger.warning("DOCX parse failed: %s", exc)
        raise ValueError("This DOCX could not be read. It may be corrupt.") from exc

    parts = [p.text for p in document.paragraphs]
    # Skills and dates frequently live in tables, which `paragraphs` skips.
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


async def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from a PDF or DOCX resume.

    Raises ValueError for an unsupported type, an oversized file, or a document
    with too little extractable text to parse.
    """
    if not file_bytes:
        raise ValueError("The uploaded file is empty.")

    if len(file_bytes) > settings.max_resume_size_bytes:
        raise ValueError(
            f"File is too large. The limit is {settings.max_resume_size_mb}MB."
        )

    extension = _extension(filename)
    if extension == ".doc":
        # Legacy OLE2, a different format from .docx despite the name. Neither
        # python-docx nor PyMuPDF reads it, and adding a converter is not worth
        # it for a format Word has not defaulted to since 2007.
        raise ValueError(
            "Legacy .doc files are not supported. Save as PDF or .docx and retry."
        )
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type {extension or '(none)'!r}. "
            f"Upload a PDF or DOCX."
        )

    raw = _extract_pdf(file_bytes) if extension == ".pdf" else _extract_docx(file_bytes)
    text = normalize_text(raw)

    if len(text) < MIN_TEXT_CHARS:
        raise ValueError(
            "Almost no text could be extracted. If this is a scanned resume, "
            "upload a version with selectable text."
        )

    logger.info("extracted %d chars from %s", len(text), filename)
    return text
